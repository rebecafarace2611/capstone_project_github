from __future__ import annotations

import json
import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
TEMPLATE = Path(
    os.environ.get(
        "RFQC_REPORT_TEMPLATE",
        Path.home()
        / "Downloads"
        / "MScBA_CapstoneReport_StudentTemplate.docx",
    )
)
ARCHIVE = ROOT / "outputs" / "rfqc" / "experiment_archive"
FINAL = ARCHIVE / "final" / "final_local_qstar_3000"
REPORT_DIR = ROOT / "reports" / "rfqc_stage_report"
FIGURE_DIR = REPORT_DIR / "figures"
OUTPUT = REPORT_DIR / "RFQC_Stage_Results_Report.docx"

BLUE = RGBColor(53, 105, 168)
DARK = RGBColor(34, 34, 34)
LIGHT_GREY = "D9D9D9"


def set_run_font(run, name: str, size: float, *, bold=None, italic=None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def configure_template_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    heading_1 = doc.styles["Heading 1"]
    heading_1.font.name = "Arial Black"
    heading_1.font.size = Pt(15)
    heading_1.font.color.rgb = DARK
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Arial Black")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Black")
    heading_1.paragraph_format.space_before = Pt(10)
    heading_1.paragraph_format.space_after = Pt(12)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = doc.styles["Heading 2"]
    heading_2.font.name = "Arial"
    heading_2.font.size = Pt(12)
    heading_2.font.bold = True
    heading_2.font.italic = True
    heading_2.font.color.rgb = DARK
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")

    for style_name in ["Caption", "Figcaption", "Tablecaption"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)
        style.font.bold = True
        style.font.color.rgb = DARK
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.0

    table_head = doc.styles["tablehead1"]
    table_head.font.name = "Arial"
    table_head.font.size = Pt(9)
    table_head.font.bold = True
    table_head._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    table_head._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    table_head.paragraph_format.space_before = Pt(2)
    table_head.paragraph_format.space_after = Pt(2)
    table_head.paragraph_format.line_spacing = 1.0

    table_text = doc.styles["tabletext1"]
    table_text.font.name = "Times New Roman"
    table_text.font.size = Pt(9)
    table_text._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    table_text._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    table_text.paragraph_format.space_before = Pt(2)
    table_text.paragraph_format.space_after = Pt(2)
    table_text.paragraph_format.line_spacing = 1.0
    table_text.paragraph_format.keep_with_next = False


def set_section_geometry(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def add_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.clear()
    run = paragraph.add_run()
    set_run_font(run, "Times New Roman", 10)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def restart_page_numbering(section, start: int = 1) -> None:
    section_properties = section._sectPr
    page_number_type = section_properties.find(qn("w:pgNumType"))
    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        section_properties.append(page_number_type)
    page_number_type.set(qn("w:start"), str(start))


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.add_run(text)


def add_heading(doc: Document, number: int, title: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.add_run(f"{number} {title}")


def add_subheading(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph(style="Heading 2")
    paragraph.add_run(title)


def add_table_caption(doc: Document, number: int, text: str) -> None:
    paragraph = doc.add_paragraph(style="Tablecaption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(f"Table {number}. {text}")


def add_figure(
    doc: Document,
    number: int,
    image_name: str,
    caption: str,
) -> None:
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.space_before = Pt(4)
    image_paragraph.paragraph_format.space_after = Pt(0)
    image_paragraph.paragraph_format.keep_with_next = True
    shape = image_paragraph.add_run().add_picture(
        str(FIGURE_DIR / image_name),
        width=Inches(5.70),
    )
    shape._inline.docPr.set("descr", caption)
    caption_paragraph = doc.add_paragraph(style="Figcaption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption_paragraph.add_run(f"Figure {number}. {caption}")


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for margin_name, margin_value in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        element = margins.find(qn(f"w:{margin_name}"))
        if element is None:
            element = OxmlElement(f"w:{margin_name}")
            margins.append(element)
        element.set(qn("w:w"), str(margin_value))
        element.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)

    for edge in ["top", "bottom"]:
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:color"), "000000")

    for edge in ["left", "right", "insideH", "insideV"]:
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "nil")

    for cell in table.rows[0].cells:
        cell_properties = cell._tc.get_or_add_tcPr()
        cell_borders = cell_properties.first_child_found_in("w:tcBorders")
        if cell_borders is None:
            cell_borders = OxmlElement("w:tcBorders")
            cell_properties.append(cell_borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "000000")
        cell_borders.append(bottom)


def set_repeat_table_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row_properties.append(table_header)


def prevent_row_split(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    row_properties.append(cant_split)


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_properties = table._tbl.tblPr
    width = table_properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        table_properties.append(width)
    total_dxa = sum(int(Cm(value).twips) for value in widths_cm)
    width.set(qn("w:w"), str(total_dxa))
    width.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_cm in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(Cm(width_cm).twips)))
        grid.append(grid_col)

    for row in table.rows:
        for index, (cell, width_cm) in enumerate(zip(row.cells, widths_cm)):
            cell.width = Cm(width_cm)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(int(Cm(width_cm).twips)))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_academic_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_cm: list[float],
    alignments: list[int] | None = None,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths_cm)
    if alignments is None:
        alignments = [WD_ALIGN_PARAGRAPH.CENTER] * len(headers)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        paragraph = cell.paragraphs[0]
        paragraph.style = doc.styles["tablehead1"]
        paragraph.alignment = alignments[index]
        run = paragraph.add_run(header)
        set_run_font(run, "Arial", 9, bold=True)

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.style = doc.styles["tabletext1"]
            paragraph.alignment = alignments[index]
            run = paragraph.add_run(str(value))
            set_run_font(run, "Times New Roman", 9)

    set_repeat_table_header(table.rows[0])
    for row in table.rows:
        prevent_row_split(row)
    set_table_widths(table, widths_cm)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)


def add_confusion_matrix_table(doc: Document, metrics: dict) -> None:
    table = doc.add_table(rows=4, cols=5)
    widths_cm = [2.5, 3.0, 3.0, 3.0, 3.0]
    set_table_widths(table, widths_cm)

    table.cell(0, 2).text = "Predicted"

    headers = ["", "", "No Fraud", "Fraud", "Class.error"]
    for index, header in enumerate(headers):
        table.cell(1, index).text = header

    table.cell(2, 0).text = "Observed"
    table.cell(2, 1).text = "No Fraud"
    table.cell(3, 0).text = ""
    table.cell(3, 1).text = "Fraud"
    table.cell(2, 2).text = f"{int(metrics['tn']):,}"
    table.cell(2, 3).text = f"{int(metrics['fp']):,}"
    table.cell(2, 4).text = f"{metrics['fpr']:.4f}"
    table.cell(3, 2).text = f"{int(metrics['fn']):,}"
    table.cell(3, 3).text = f"{int(metrics['tp']):,}"
    table.cell(3, 4).text = f"{metrics['fnr']:.4f}"

    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if row_index in [0, 1]:
                paragraph.style = doc.styles["tablehead1"]
                for run in paragraph.runs:
                    set_run_font(run, "Arial", 9, bold=True)
            else:
                paragraph.style = doc.styles["tabletext1"]
                for run in paragraph.runs:
                    set_run_font(run, "Times New Roman", 9)

    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])
    set_table_widths(table, widths_cm)
    set_table_borders(table)

    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    inside_horizontal = borders.find(qn("w:insideH"))
    inside_horizontal.set(qn("w:val"), "nil")

    for cell in table.rows[0].cells:
        cell_properties = cell._tc.get_or_add_tcPr()
        cell_borders = cell_properties.first_child_found_in("w:tcBorders")
        if cell_borders is None:
            cell_borders = OxmlElement("w:tcBorders")
            cell_properties.append(cell_borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "nil")
        cell_borders.append(bottom)

    for cell in table.rows[1].cells:
        cell_properties = cell._tc.get_or_add_tcPr()
        cell_borders = cell_properties.first_child_found_in("w:tcBorders")
        if cell_borders is None:
            cell_borders = OxmlElement("w:tcBorders")
            cell_properties.append(cell_borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:color"), "000000")
        cell_borders.append(bottom)

    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)


def pct(value: float, digits: int = 2) -> str:
    return f"{value * 100:.{digits}f}%"


def build_report() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(TEMPLATE)
    clear_document_body(doc)
    configure_template_styles(doc)

    for section in doc.sections:
        set_section_geometry(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    title.paragraph_format.space_after = Pt(24)
    run = title.add_run("Random Forest Quantile Classifier (RFQC)")
    set_run_font(run, "Times New Roman", 20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Stage Results Report")
    set_run_font(run, "Times New Roman", 16, bold=True)

    programme = doc.add_paragraph()
    programme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    programme.paragraph_format.space_before = Pt(48)
    programme.paragraph_format.space_after = Pt(12)
    run = programme.add_run("MSc in Business Analytics Capstone Project")
    set_run_font(run, "Times New Roman", 12)

    institution = doc.add_paragraph()
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    institution.paragraph_format.space_after = Pt(72)
    run = institution.add_run(
        "Michael Smurfit Graduate School of Business, University College Dublin"
    )
    set_run_font(run, "Times New Roman", 12)

    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_paragraph.add_run("13 June 2026")
    set_run_font(run, "Times New Roman", 12)

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_geometry(main_section)
    main_section.header.is_linked_to_previous = False
    main_section.footer.is_linked_to_previous = False
    restart_page_numbering(main_section, 1)
    add_page_number(main_section)

    config = json.loads((FINAL / "best_configuration.json").read_text(encoding="utf-8"))
    test_metrics = json.loads(
        (FINAL / "final_test_metrics.json").read_text(encoding="utf-8")
    )["primary_test_metrics"]
    results = pd.read_csv(FINAL / "rfqc_results_summary.csv")
    convergence = pd.read_csv(FINAL / "final_tree_convergence.csv")
    threshold_curve = pd.read_csv(FINAL / "final_threshold_curve.csv")
    baseline = pd.read_csv(ARCHIVE / "baseline_summary.csv")
    local = pd.read_csv(
        ARCHIVE / "tuning" / "02_local_refine_gini_500" / "cv_ranking.csv"
    )
    quick = pd.read_csv(
        ARCHIVE / "tuning" / "01_quick_search_500" / "cv_ranking.csv"
    )

    cv_locked = results[
        (results["stage"] == "cross_validation")
        & (results["threshold_rule"] == "q_star_prevalence")
    ].iloc[0]
    cv_optimised = results[
        (results["stage"] == "cross_validation")
        & (results["threshold_rule"] == "gmean_optimized")
    ].iloc[0]
    final_oob = convergence.iloc[-1]
    baseline_3000 = baseline[baseline["ntree"] == 3000].iloc[0]
    selected_threshold = threshold_curve[
        threshold_curve["is_selected_threshold"]
    ].iloc[0]
    maximum_threshold = threshold_curve.loc[threshold_curve["gmean"].idxmax()]

    add_heading(doc, 1, "Executive Summary")
    add_body_paragraph(
        doc,
        "This report summarises the current RFQC modelling stage for fraud "
        "classification. The final model was trained on 444,074 observations, "
        "including 1,860 fraud cases, using 144 approved predictors. Model structure "
        "and the classification threshold were selected using the training data and "
        "five-fold cross-validation before the final test set was evaluated.",
    )
    add_body_paragraph(
        doc,
        "The locked model used 3,000 trees, mtry = 24, terminal node size = 20, "
        "nsplit = 10 and Gini splitting. The threshold was fixed at the training fraud "
        f"prevalence (q* = {config['locked_threshold']:.8f}). Five-fold validation "
        f"produced a mean G-mean of {cv_locked['gmean']:.4f} "
        f"(SD {config['cv_metrics']['std_validation_gmean']:.4f}), while the untouched "
        f"test G-mean was {test_metrics['gmean']:.4f}. Test ROC-AUC was "
        f"{test_metrics['roc_auc']:.4f} and PR-AUC was "
        f"{test_metrics['pr_auc']:.4f}.",
    )
    add_body_paragraph(
        doc,
        f"At the locked threshold, the model identified {test_metrics['tp']:.0f} of "
        f"465 fraud cases ({pct(test_metrics['sensitivity'])} sensitivity) and retained "
        f"{pct(test_metrics['specificity'])} specificity. The resulting precision was "
        f"{pct(test_metrics['precision'])}, reflecting both the low fraud prevalence "
        f"({pct(465 / 111020, 3)}) and an intentionally sensitivity-oriented operating "
        "point. The close agreement between cross-validation and test discrimination "
        "indicates that the selected RFQC generalised without a material performance "
        "collapse.",
    )

    add_heading(doc, 2, "RFQC Experimental Protocol")
    add_body_paragraph(
        doc,
        "The implementation used the native random forest quantile classifier in "
        "randomForestSRC. A two-stage structural search was conducted with 500-tree "
        "forests. The quick search evaluated eight combinations of split rule, mtry "
        "and terminal node size. A four-candidate local search then refined the best "
        "Gini structures. Each candidate was evaluated using the same five folds and "
        "random seed (42).",
    )
    add_body_paragraph(
        doc,
        "After structural selection, the model was refitted on the complete training "
        "set with 3,000 trees. The q* threshold was defined only from the training "
        "fraud prevalence and was locked with the model configuration. The test set "
        "was not used to select parameters or revise the threshold.",
    )
    add_body_paragraph(
        doc,
        "PR-AUC is retained as the reporting label used throughout the experiment "
        "archive. Computationally, it was evaluated using average precision, namely "
        "the mean precision at the ranked positions of the positive cases.",
    )

    add_table_caption(doc, 1, "RFQC data and locked model specification.")
    add_academic_table(
        doc,
        ["Item", "Specification"],
        [
            ["Training data", "444,074 records; 1,860 fraud cases"],
            ["Final test data", "111,020 records; 465 fraud cases"],
            ["Predictor set", "144 model features"],
            ["Implementation", "R randomForestSRC native RFQ"],
            ["Locked forest", "ntree 3,000; mtry 24; node size 20; nsplit 10"],
            ["Split rule", "Gini"],
            ["Threshold rule", f"q* training prevalence ({config['locked_threshold']:.8f})"],
            ["Random seed", "42"],
        ],
        [4.1, 10.4],
        [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body_paragraph(
        doc,
        "The first test-file read stopped during strict validation of previously "
        "unseen factor levels, before any prediction or metric was produced. After "
        "the compatibility handling was corrected, the same saved model and the same "
        "locked threshold were used for evaluation. No parameter or threshold was "
        "changed following test access.",
    )

    doc.add_page_break()
    add_heading(doc, 3, "Hyperparameter Tuning")
    quick_locked = quick[
        (quick["candidate"] == 4)
        & (quick["threshold_rule"] == "q_star_prevalence")
    ].iloc[0]
    local_locked = local[
        (local["candidate"] == 3)
        & (local["threshold_rule"] == "q_star_prevalence")
    ].iloc[0]
    local_opt = local[
        (local["candidate"] == 3)
        & (local["threshold_rule"] == "gmean_optimized")
    ].iloc[0]

    add_subheading(doc, "Cross-validation design")
    add_body_paragraph(
        doc,
        "Model selection used five-fold stratified group cross-validation on the "
        "444,074-row training set. Each validation fold contained 88,814 to 88,816 "
        "records and exactly 372 fraud cases. Grouping was based on the complete "
        "approved feature vector, so identical predictor profiles could not be split "
        "between training and validation folds.",
    )
    add_body_paragraph(
        doc,
        "For each candidate and fold, four folds were used to fit the RFQC and the "
        "remaining fold was used only for validation. The maximum-G-mean and q* "
        "thresholds were calculated from the fitted model's training-fold OOB "
        "predictions and then applied unchanged to the held-out fold. Candidate "
        "ranking was based primarily on mean validation G-mean across the five folds; "
        "fold standard deviation and sensitivity-specificity imbalance were used to "
        "assess stability and operating balance.",
    )

    add_subheading(doc, "Cross-validation results")
    add_body_paragraph(
        doc,
        "The quick search favoured Gini splitting with mtry = 24 and terminal node "
        f"size = 10, giving a q* validation G-mean of {quick_locked['mean_validation_gmean']:.4f}. "
        "Local refinement showed that increasing the terminal node size to 20 improved "
        "the structural result. The selected structure achieved similar performance "
        "with either threshold rule.",
    )
    add_table_caption(doc, 2, "Cross-validation results used for model selection.")
    add_academic_table(
        doc,
        [
            "Search and configuration",
            "Threshold",
            "G-mean",
            "SD",
            "Sensitivity",
            "Specificity",
            "Balance gap",
        ],
        [
            [
                "Quick: Gini, mtry 24, node 10",
                "q*",
                f"{quick_locked['mean_validation_gmean']:.4f}",
                f"{quick_locked['std_validation_gmean']:.4f}",
                f"{quick_locked['mean_sensitivity']:.4f}",
                f"{quick_locked['mean_specificity']:.4f}",
                f"{quick_locked['mean_balance_gap']:.4f}",
            ],
            [
                "Local: Gini, mtry 24, node 20",
                "Max G-mean",
                f"{local_opt['mean_validation_gmean']:.4f}",
                f"{local_opt['std_validation_gmean']:.4f}",
                f"{local_opt['mean_sensitivity']:.4f}",
                f"{local_opt['mean_specificity']:.4f}",
                f"{local_opt['mean_balance_gap']:.4f}",
            ],
            [
                "Local: Gini, mtry 24, node 20",
                "q* (locked)",
                f"{local_locked['mean_validation_gmean']:.4f}",
                f"{local_locked['std_validation_gmean']:.4f}",
                f"{local_locked['mean_sensitivity']:.4f}",
                f"{local_locked['mean_specificity']:.4f}",
                f"{local_locked['mean_balance_gap']:.4f}",
            ],
        ],
        [4.5, 2.0, 1.55, 1.25, 1.75, 1.75, 1.7],
        [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )
    add_body_paragraph(
        doc,
        f"The maximum-G-mean threshold exceeded the q* result by only "
        f"{local_opt['mean_validation_gmean'] - local_locked['mean_validation_gmean']:.4f} "
        "in mean validation G-mean. The q* rule reduced the mean validation false-positive "
        f"rate from {pct(1 - local_opt['mean_specificity'])} to "
        f"{pct(1 - local_locked['mean_specificity'])} and reduced the sensitivity-specificity "
        f"gap from {local_opt['mean_balance_gap']:.4f} to "
        f"{local_locked['mean_balance_gap']:.4f}. It was therefore locked as the more "
        "balanced and pre-specifiable operating rule.",
    )
    add_body_paragraph(
        doc,
        "For the locked structure and q* rule, fold-level validation G-mean ranged "
        "from 0.7771 to 0.7969 (mean 0.7862, SD 0.0094). Mean sensitivity was 0.8199 "
        "(SD 0.0191), whereas mean specificity was 0.7541 (SD 0.0021). ROC-AUC was "
        "0.8688 (SD 0.0064) and PR-AUC was 0.0283 (SD 0.0008). The relatively small "
        "dispersion, particularly for specificity and PR-AUC, indicated that the "
        "selected configuration was not dependent on a single favourable fold.",
    )
    add_figure(
        doc,
        1,
        "figure_1_model_selection.png",
        "Structural model selection during five-fold cross-validation. Mean validation "
        "G-mean is shown for the q* prevalence threshold during (a, b) the quick search "
        "and (c) local refinement. Cell text reports the fold mean and standard "
        "deviation; the outlined cell marks the locked structure.",
    )

    add_heading(doc, 4, "Final Forest Size and Tree Stability")
    add_body_paragraph(
        doc,
        "After Section 3 selected the structural parameters (split rule, mtry and "
        "terminal node size), that fixed structure was refitted at increasing forest "
        "sizes to select the final number of trees and assess out-of-bag stability. "
        "G-mean increased from 0.7826 at 500 trees to 0.7880 at 2,000 trees and changed "
        "only slightly at 3,000 trees. ROC-AUC and PR-AUC also showed "
        "diminishing improvements. These results support 3,000 trees as a stable final fit rather "
        "than evidence that a larger forest would materially alter operating performance.",
    )
    add_table_caption(doc, 3, "Out-of-bag convergence for the selected RFQC structure.")
    add_academic_table(
        doc,
        ["Trees", "G-mean", "Sensitivity", "Specificity", "ROC-AUC", "PR-AUC"],
        [
            [
                f"{int(row.trees):,}",
                f"{row.gmean:.4f}",
                f"{row.sensitivity:.4f}",
                f"{row.specificity:.4f}",
                f"{row.roc_auc:.4f}",
                f"{row.pr_auc:.4f}",
            ]
            for row in convergence.itertuples()
        ],
        [2.0, 2.3, 2.55, 2.55, 2.55, 2.55],
    )
    add_body_paragraph(
        doc,
        f"At 3,000 trees, the tuned structure achieved an OOB q* G-mean of "
        f"{final_oob['gmean']:.4f}, compared with {baseline_3000['q_star_gmean']:.4f} "
        "for the untuned 3,000-tree baseline. The improvement of "
        f"{final_oob['gmean'] - baseline_3000['q_star_gmean']:.4f} supports the value "
        "of the structural search. The 3,000-tree PR-AUC in Table 3 is the training OOB "
        f"value ({final_oob['pr_auc']:.4f}); it is not the final test PR-AUC "
        f"({test_metrics['pr_auc']:.4f}).",
    )

    doc.add_page_break()
    add_heading(doc, 5, "Threshold Selection")
    add_body_paragraph(
        doc,
        "The final training OOB threshold curve shows the expected sensitivity-specificity "
        "trade-off. Because FPR equals 1 minus specificity, a separate FPR curve would "
        "duplicate the specificity curve; FPR is therefore reported explicitly in the "
        "text. The maximum OOB G-mean occurred at a threshold of "
        f"{maximum_threshold['threshold']:.8f}, with sensitivity "
        f"{maximum_threshold['sensitivity']:.4f}, specificity "
        f"{maximum_threshold['specificity']:.4f} and G-mean "
        f"{maximum_threshold['gmean']:.4f}.",
    )
    add_body_paragraph(
        doc,
        f"The locked q* threshold ({selected_threshold['threshold']:.8f}) produced "
        f"sensitivity {selected_threshold['sensitivity']:.4f}, specificity "
        f"{selected_threshold['specificity']:.4f}, FPR "
        f"{selected_threshold['fpr']:.4f} and G-mean "
        f"{selected_threshold['gmean']:.4f}. The maximum-G-mean threshold had an FPR of "
        f"{maximum_threshold['fpr']:.4f}. Therefore, q* reduced the training OOB FPR by "
        f"{maximum_threshold['fpr'] - selected_threshold['fpr']:.4f} "
        f"({(maximum_threshold['fpr'] - selected_threshold['fpr']) * 100:.2f} percentage "
        f"points) while sacrificing only "
        f"{maximum_threshold['gmean'] - selected_threshold['gmean']:.4f} G-mean. It also "
        "reduced the sensitivity-specificity gap from "
        f"{maximum_threshold['balance_gap']:.4f} to "
        f"{selected_threshold['balance_gap']:.4f}. This supported retaining q* without "
        "using the test set to optimise the operating point.",
    )
    add_figure(
        doc,
        2,
        "figure_2_threshold_selection.png",
        "Training out-of-bag performance across candidate classification thresholds. "
        "The filled square and blue vertical line identify the locked q* prevalence "
        "threshold. The open circle and grey line identify the threshold with maximum "
        "training OOB G-mean.",
    )

    add_heading(doc, 6, "Final Test Performance")
    add_body_paragraph(
        doc,
        "The locked model and threshold were applied once to the final test evaluation "
        "after factor-level compatibility was corrected. Test G-mean was 0.7808, only "
        f"{cv_locked['gmean'] - test_metrics['gmean']:.4f} below the five-fold validation "
        "mean. Test ROC-AUC differed from cross-validation by "
        f"{abs(cv_locked['roc_auc'] - test_metrics['roc_auc']):.4f}. Specificity was "
        "essentially unchanged, while sensitivity was modestly lower on the test set.",
    )
    add_table_caption(
        doc,
        4,
        "Performance comparison and final test classification counts.",
    )
    panel_a = doc.add_paragraph()
    panel_a.paragraph_format.space_after = Pt(3)
    panel_a.paragraph_format.keep_with_next = True
    run = panel_a.add_run("Panel A. Performance across evaluation stages")
    set_run_font(run, "Times New Roman", 9, bold=True)
    add_academic_table(
        doc,
        ["Evaluation", "G-mean", "Sensitivity", "Specificity", "FPR", "ROC-AUC", "PR-AUC"],
        [
            [
                "Untuned baseline OOB (3,000 trees)",
                f"{baseline_3000['q_star_gmean']:.4f}",
                f"{baseline_3000['q_star_sensitivity']:.4f}",
                f"{baseline_3000['q_star_specificity']:.4f}",
                f"{baseline_3000['q_star_fpr']:.4f}",
                f"{baseline_3000['roc_auc']:.4f}",
                f"{baseline_3000['pr_auc']:.4f}",
            ],
            [
                "Selected five-fold CV mean",
                f"{cv_locked['gmean']:.4f}",
                f"{cv_locked['sensitivity']:.4f}",
                f"{cv_locked['specificity']:.4f}",
                f"{cv_locked['fpr']:.4f}",
                f"{cv_locked['roc_auc']:.4f}",
                f"{cv_locked['pr_auc']:.4f}",
            ],
            [
                "Final training OOB (3,000 trees)",
                f"{final_oob['gmean']:.4f}",
                f"{final_oob['sensitivity']:.4f}",
                f"{final_oob['specificity']:.4f}",
                f"{final_oob['fpr']:.4f}",
                f"{final_oob['roc_auc']:.4f}",
                f"{final_oob['pr_auc']:.4f}",
            ],
            [
                "Untouched final test",
                f"{test_metrics['gmean']:.4f}",
                f"{test_metrics['sensitivity']:.4f}",
                f"{test_metrics['specificity']:.4f}",
                f"{test_metrics['fpr']:.4f}",
                f"{test_metrics['roc_auc']:.4f}",
                f"{test_metrics['pr_auc']:.4f}",
            ],
        ],
        [4.6, 1.65, 1.65, 1.65, 1.65, 1.65, 1.65],
        [
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.CENTER,
        ],
    )
    panel_b = doc.add_paragraph()
    panel_b.paragraph_format.space_before = Pt(2)
    panel_b.paragraph_format.space_after = Pt(3)
    panel_b.paragraph_format.keep_with_next = True
    run = panel_b.add_run("Panel B. Final test confusion matrix")
    set_run_font(run, "Times New Roman", 9, bold=True)
    add_confusion_matrix_table(doc, test_metrics)
    matrix_note = doc.add_paragraph()
    matrix_note.paragraph_format.space_before = Pt(2)
    matrix_note.paragraph_format.space_after = Pt(6)
    matrix_note.paragraph_format.line_spacing = 1.0
    run = matrix_note.add_run(
        "Note: Class error is the false-positive rate for the observed No Fraud "
        "class [26,986 / (83,569 + 26,986) = 0.2441] and the false-negative rate "
        "for the observed Fraud class [90 / (90 + 375) = 0.1935]."
    )
    set_run_font(run, "Times New Roman", 9, italic=True)
    add_body_paragraph(
        doc,
        f"The model detected {int(test_metrics['tp'])} fraud cases and missed "
        f"{int(test_metrics['fn'])}. It generated {int(test_metrics['fp']):,} false "
        f"positive alerts from {int(test_metrics['tn'] + test_metrics['fp']):,} "
        f"non-fraud cases, corresponding to a test FPR of "
        f"{pct(test_metrics['fpr'])}. Precision was "
        f"{pct(test_metrics['precision'])}. "
        "This low precision does not contradict the ROC-AUC result: under extreme class "
        "imbalance, a moderate false-positive rate produces many more false alerts than "
        "true fraud detections. The precision-recall curve and PR-AUC therefore provide "
        "the more informative view of positive-class performance.",
    )
    add_figure(
        doc,
        3,
        "figure_3_test_discrimination.png",
        "Discrimination performance on the untouched final test set. (a) Receiver "
        "operating characteristic curve and (b) precision-recall curve. Filled squares "
        "show the operating point produced by the pre-locked q* threshold. The dashed "
        "horizontal line in panel (b) is the test fraud prevalence.",
    )
    add_body_paragraph(
        doc,
        f"The final operating point achieved sensitivity {test_metrics['sensitivity']:.4f}, "
        f"specificity {test_metrics['specificity']:.4f}, G-mean "
        f"{test_metrics['gmean']:.4f}, ROC-AUC {test_metrics['roc_auc']:.4f} and "
        f"PR-AUC {test_metrics['pr_auc']:.4f}. These results characterise a model with "
        "strong ranking discrimination and high fraud capture, but with a substantial "
        "alert burden at the selected balanced threshold.",
    )

    doc.core_properties.title = "RFQC Stage Results Report"
    doc.core_properties.subject = "MSc Business Analytics Capstone Project"
    doc.core_properties.author = "Capstone Project Team"
    doc.core_properties.keywords = "RFQC, fraud detection, random forest, stage report"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
